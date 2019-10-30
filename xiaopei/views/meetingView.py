# -*- coding: utf-8 -*-
 
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.shortcuts import render
from rest_framework import status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from datetime import datetime
import pageView as pview
from xiaopei.views import run
from django.core import serializers
from ..models import User, Meeting
from django.http import FileResponse
from pydub import AudioSegment
import wave
from docx import Document
import socket
import docx
import struct
import io
import pyaudio
import json
import datetime
import os
import speech_separation_energy as speechSep
import time

serverIp = '10.3.27.97'; #语音转写模块网络ip地址
serverPort = 8086; #语音转写模块网络端口
addr = (serverIp, serverPort);
project_path = os.getcwd()
sound_silence = AudioSegment.from_file(project_path + '/silence.wav', format='wav')
silence_2s = sound_silence[1000:2000]  # 取前(n/1000)秒音频
number = 0
name = ''
ltime = ''

def getUser(userName): #根据用户名查询用户
	try:
		user = User.objects.get(userName=userName); #根据用户名查询用户
		if (user is None):
			return {};
		return user;
	except Exception, e:
		return {};

def getMeeting(id): #根据用户id查询会议
	try:
		meeting = Meeting.objects.get(id=id); #根据用户id查询会议
		if (meeting is None):
			return {};
		return meeting;
	except Exception, e:
		return {};

def saveMeeting(userName, meetingName, description, resultType): #保存会议
	meeting = Meeting();
	meeting.meetingName = meetingName; #会议名称
	meeting.description = description; #会议描述
	meeting.resultType = resultType; #会议记录种类，txt or word
	meeting.userName = userName; #用户名
	meeting.status = "ready"; #会议状态
	meeting.language = u"普通话"; #语言
	meeting.audioType = "wav"; #音频种类
	meeting.audioBit =16; #音频位数，已废弃
	meeting.audioFrequency = 8; #音频采样频率，已废弃
	meeting.save(); #保存会议
	return meeting;

def createMeeting(request): #创建实时音频会议
	userName = request.session.get('userName', None);
	if userName is None:
		return HttpResponseRedirect('/');
	meetingName = request.POST.get('meetingName');
	description = request.POST.get('description');
	resultType = request.POST.get('resultType');
	meeting = saveMeeting(userName, meetingName, description, resultType); #创建会议
	meeting.wavname = str(meeting.id); #初始化音频文件名
	meeting.txtname = str(meeting.id); #初始化文本文件名
	meeting.resultUrl = "none";
	meeting.audioUrl = "none";
	meeting.save();
	data = {}; #返回数据格式
	data['meetings'] = Meeting.objects.all(); #返回会议列表
	data['activeUser'] = getUser(userName); #返回用户用户信息
	return pview.getMeetingPage(request, meeting.id);
	#return render(request, 'meetings.html', data);

def createUploadMeeting(request): #创建非实时音频会议
	userName = request.session.get('userName', None);
	if userName is None:
		return HttpResponseRedirect('/');
	if request.method == 'POST':
		files = request.FILES.getlist('file'); #获取音频文件
		meetingName = request.POST.get('meetingName');
		description = request.POST.get('description');
		resultType = request.POST.get('resultType');
		file = files[0]; #获取音频文件
		f = wave.open(file,"rb")
		baseDir = os.path.dirname(os.path.abspath(__name__));
		wavdir = os.path.join(baseDir,'static','wav'); #音频文件地址
		txtdir = os.path.join(baseDir,'static','txt'); #文本文件地址
		meeting = saveMeeting(userName, meetingName, description, resultType); #保存会议
		(id) = meeting.id;
		id = str(id);
		fwav = open(wavdir + "/" + str(id) + ".wav",'wb');
		for chrunk in file.chunks():
			fwav.write(chrunk); #保存会议音频
		fwav.close(); 
		finput = wavdir + "/" + str(id) + ".wav";
		text = u"";
		sound = AudioSegment.from_file(file, format="wav") 
		sound = sound.set_sample_width(4); #设置32位
		sound = sound.set_frame_rate(8000); #设置8KHZ
		sound = sound.set_channels(1); #设置单声道
		soundLen = len(sound) #获取音频时长
		i = 0; d=10000;
		while i<soundLen: #截取音频
			sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM); #建立sockect
			sock.connect(addr);
			record = sound[i:min(i + d, soundLen)];
			send = struct.pack('>i', len(record.raw_data)) + record.raw_data; #创建音频数据包
			sock.send(send); #发送
			result = sock.recv(1024); #接收
			# result = "你好";
			text = text + result.decode('utf-8');
			i = i + d;
			sock.close(); #关闭socket
		if resultType == "txt" : #保存txt文件
			ftxt = io.open(txtdir + "/" + str(id) + ".txt",'w',encoding='UTF-8');
			ftxt.write(text);
			ftxt.close();
		else : #保存word文件
			docdir = os.path.join(baseDir,'static','doc');
			document = Document();
			document.add_paragraph(text)
			document.save(docdir + "/" + str(id) + '.docx');
		meeting.audioUrl = "/download/wav/" + str(id) + "/"
		meeting.resultUrl = "/download/txt/" + str(id) + "/";
		meeting.wavname = str(id); #初始化音频文件名
		meeting.txtname = str(id); #初始化文本文件名
		meeting.status = "finish";
		meeting.wavdate = datetime.datetime.now();
		meeting.txtdate = datetime.datetime.now();
		meeting.save(); #保存会议
	return pview.getMeetingPage(request, id);


@api_view(['POST'])
def openUploadMeetings(request): #创建非实时音频会议的restful接口
	global number, silence_2s
	userName = request.POST.get('userName');
	meetingName = request.POST.get('meetingName');
	description = request.POST.get('description');
	resultType = request.POST.get('resultType');
	files = request.FILES.getlist('file')
	file = files[0];
	try:
		f = wave.open(file, "rb") #获取音频文件
	except Exception,e:
		result = {};
		result['statusCode'] = 400;
		result["message"] = "音频编码格式错误";
		return Response(result,status=status.HTTP_200_OK);
	baseDir = os.path.dirname(os.path.abspath(__name__));
	wavdir = os.path.join(baseDir,'static','wav'); #获取音频文件路径
	txtdir = os.path.join(baseDir,'static','txt'); #获取文本文件保存路径
	meeting = saveMeeting(userName, meetingName, description, resultType); #保存会议
	(id) = meeting.id;
	id = str(id);
	fwav = open(wavdir + "/" + id + ".wav",'wb');
	for chrunk in file.chunks():
		fwav.write(chrunk); #保存音频文件
	fwav.close();
	text = u"";
	sound = AudioSegment.from_file(file, format="wav")
	sound = sound.set_sample_width(4); #32位
	sound = sound.set_frame_rate(8000);#8khz
	sound = sound.set_channels(1);#单声道
	'''soundLen = len(sound)
	i = 0; d=10000;
	while i<soundLen:
		sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM); #建立sockect
		sock.connect(addr); #建立socket连接
		record = sound[i:min(i + d, soundLen)]; #截取音频
		send = struct.pack('>i', len(record.raw_data)) + record.raw_data; #打包数据
		sock.send(send); #发送信息
		result = sock.recv(1024); #接受返回结果
		#result = "你好";
		text = text + result.decode('utf-8'); #拼接翻译结果
		i = i + d;
		sock.close();'''
	# 将非实时音频分割程序 import进来，将原程序改为一个函数，这个函数将原始连续音频数据 return 切分好的音频序列
	# 注意：将返回的音频数据数组sample_arrays转换为音频对象audioSegment
	splited_speech_list = speechSep.speech_separate_notRealTime(sound);
	print len(splited_speech_list)
    	for audio_samples in splited_speech_list:
		# 判断是否出现换行标志return
		if audio_samples == 'return':
			text = text[:-1] + u'。' + '\n';
			continue
		#判断是否出现silence_2s更新标志，字典dict
		if isinstance(audio_samples, dict):
			# silence_2s = sound._spawn(audio_samples['silence_2s'])
			continue
		#convert samples arrays to audioSegment object
		number += 1
		print number
		audio = sound._spawn(audio_samples)
		audio = silence_2s + audio + silence_2s
		# wav_audio = audio.set_sample_width(4);  # 32位
		# wav_audio = wav_audio.set_frame_rate(8000);  # 8khz
		# wav_audio = wav_audio.set_channels(1);  # 单声道
		# wav_audio.export(project_path + '/' + str(number) + '.wav', format='wav')
		sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM);
		sock.connect(addr);
		record = audio;
		send = struct.pack('>i', len(record.raw_data)) + record.raw_data;
		sock.send(send);
		result = sock.recv(len(audio));
		# result = "你好";
		text = text + result.decode('utf-8') + u'，'
		sock.close();
		time.sleep(1)
	text = text[:-1] + u'。'
	if resultType == "txt" : #txt文件保存
		ftxt = io.open(txtdir + "/" + str(id) + ".txt",'w',encoding='UTF-8'); #写入结果
		ftxt.write(text);
		ftxt.close();
	else : #word文件保存
		docdir = os.path.join(baseDir,'static','doc');
		document = Document();
		document.add_paragraph(text)
		document.save(docdir + "/" + str(id) + '.docx');
	meeting.audioUrl = "/download/wav/" + str(id) + "/"
	meeting.resultUrl = "/download/txt/" + str(id) + "/";
	meeting.wavname = str(id);
	meeting.txtname = str(id);
	meeting.status = "finish";
	meeting.wavdate = datetime.datetime.now();
	meeting.txtdate = datetime.datetime.now();
	meeting.save(); #保存会议
	result = {};
	result['statusCode'] = 200;
	result["message"] = "ok";
	result["meeting"] = getMeetingObject(meeting);
	return Response(result,status=status.HTTP_200_OK);

def deleteMeeting(request, id):#删除会议
	userName = request.session.get('userName', None);
	if userName is None:
		return HttpResponseRedirect('/');
	try:
		Meeting.objects.filter(id=id).delete();
	except Exception, e:
		print e;
	return pview.getMeetingsPage(request);

@api_view(['POST'])
def recordAnalysis(request):#实时音频处理
	global number
	result = {};
	if request.method == 'POST':
		number = number + 1
		files = request.FILES.getlist('file')
		file = files[0];
		addr = (serverIp, serverPort);
		sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM);
		sock.connect(addr);
		wav_audio = AudioSegment.from_file(file, format="wav")
		# strength the loud of  audio, ndb.
		wav_audio = silence_2s + wav_audio + silence_2s
		wav_audio = wav_audio.set_sample_width(4); #32位
		wav_audio = wav_audio.set_frame_rate(8000); #8khz
		wav_audio = wav_audio.set_channels(1); #单声道
		wav_audio.export(project_path + '/' + str(number) + '.wav', format= 'wav')
		baseDir = os.path.dirname(os.path.abspath(__name__));
		wavdir = os.path.join(baseDir,'static','wav');
		message = wav_audio.raw_data
		send = struct.pack('>i', len(wav_audio.raw_data)) + wav_audio.raw_data;
		sock.send(send); #发送信息
		trans = sock.recv(1024); #接受翻译结果
		try:
			trans = run.change(trans);
		except:
			pass;
		sock.close();
		# trans = "你好"
		result['statusCode'] = 200;
		result["message"] = "ok";
		if name == '' :
			result['name'] = '您'
		else :
			result['name'] = name
		print trans
		print result['name']
		if trans == "" or trans == "None":
			result["result"] = "";
		else :
			result["result"] = trans;
		#return HttpResponse(getRecordMessage(file.read(1024)));
	return Response(result,status=status.HTTP_200_OK);

@api_view(['POST'])
def nameMatch(request):
    result = {}
    requestData = json.loads(request.body)
    item = requestData['name']
    global name,ltime
    if name == item:
        result['statusCode'] = 200;
        result["message"] = "ok";
        result['name'] = name
        # print name
    else:
        name = item
        result['statusCode'] = 200;
        result["message"] = "ok";
        result['name'] = name
        # print name
    return Response(result, status=status.HTTP_200_OK);

@api_view(['POST'])
def saveData(request): #保存energy and zero数据到本地txt文件
	requestData = json.loads(request.body)
	energy = requestData['energy']
	print energy
	with open(project_path + "/energy.txt", 'w') as energy_output:
		energy_output.writelines(str(energy))
	zero = requestData['zero']
	print zero
	with open(project_path + "/zero.txt", 'w') as zero_output:
		zero_output.writelines(str(zero))
	result = '你好'
	return Response(result, status=status.HTTP_200_OK);


@api_view(['POST'])
def updateSilence(request): #换行时更新silence_2s
	pass
	result = {}
	# global silence_2s
	# files = request.FILES.getlist('return_silence')
	# file = files[0]
	# wav_audio = AudioSegment.from_file(file, format= 'wav')
	# wav_audio = wav_audio.set_sample_width(4);  # 32位
	# wav_audio = wav_audio.set_frame_rate(8000);  # 8khz
	# wav_audio = wav_audio.set_channels(1);  # 单声道
	# silence_2s = wav_audio
	result['result'] = '你好'
	return Response(result, status=status.HTTP_200_OK);

def recordReplace(request, id): #替换音频文件
	try:
		meeting = Meeting.objects.get(id=id); #获取会议
	except Exception, e:
		return HttpResponse("error id");
	userName = request.session.get('userName', None);
	if userName is None:
		return HttpResponseRedirect('/');
	if request.method == 'POST':
		files = request.FILES.getlist('file') #获取音频文件
		file = files[0];
		f = wave.open(file,"rb")
		baseDir = os.path.dirname(os.path.abspath(__name__));
		wavdir = os.path.join(baseDir,'static','wav');#获取音频文件地址
		fwav = open(wavdir + "/" + str(id) + ".wav",'wb');
		for chrunk in file.chunks():
			fwav.write(chrunk);
		fwav.close();
		meeting.status = "finish"; #变更会议状态
		meeting.audioUrl = "/download/wav/" + str(id) + "/"
		meeting.wavdate = datetime.datetime.now();
		meeting.save();
	return pview.getMeetingPage(request, id);

def recordTxtReplace(request, id):#txt 文件替换
	try:
		meeting = Meeting.objects.get(id=id);#获取会议
	except Exception, e:
		return HttpResponse("error id");
	if request.method == 'POST':
		files = request.FILES.getlist('txtFile')#获取文本文件
		file = files[0];
		fileType = ((file.name).split("."))[-1];
		baseDir = os.path.dirname(os.path.abspath(__name__));
		if (fileType == "txt"):#读取文本文件
			txtdir = os.path.join(baseDir,'static','txt');
			ftxt = io.open(txtdir + "/" + str(id) + ".txt",'wb');
			line = file.readline();
			while line:
				#print type(line);
				#print line;
				ftxt.write(line);
				line = file.readline();
			meeting.resultType = "txt";
			ftxt.close();
		else :#读取word文件
			document = docx.Document(file)
			resultD = '\n\n'.join([paragraph.text.encode('utf-8') for paragraph in document.paragraphs])
			docdir = os.path.join(baseDir,'static','doc');
			document = Document();
			document.add_paragraph(resultD.decode("utf-8"))
			document.save(docdir + "/" + str(id) + '.docx');
			meeting.resultType = "doc";
		meeting.status = "finish";
		meeting.resultUrl = "/download/txt/" + str(id) + "/"
		meeting.txtdate = datetime.datetime.now();
		meeting.save();
	return pview.getMeetingPage(request, id);

@api_view(['POST'])
def uploadRecord(request, id): #上传音频文本 结束实时会议
	result = {};
	try:
		meeting = Meeting.objects.get(id=id);#获取会议
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
			return Response(result,status=status.HTTP_200_OK);
	except Exception, e:
		result['statusCode'] = 404;
		result["message"] = "not find";
		return Response(result,status=status.HTTP_200_OK);
	if request.method == 'POST':
		file = (request.FILES.getlist('file'))[0];
		text = request.POST.get('text');

		baseDir = os.path.dirname(os.path.abspath(__name__));#文件路径
		wavdir = os.path.join(baseDir,'static','wav');
		txtdir = os.path.join(baseDir,'static','txt');
		docdir = os.path.join(baseDir,'static','doc');
		fwav = open(wavdir + "/" + str(id) + ".wav",'wb'); #获取音频文件
		for chrunk in file.chunks():
			fwav.write(chrunk);
		fwav.close();

		meeting.audioUrl = "/download/wav/" + str(id) + "/"
		if meeting.resultType == "txt": #文本文件
			ftxt = io.open(txtdir + "/" + str(id) + ".txt",'w',encoding='UTF-8');
			ftxt.write(text);
			ftxt.close();
		else :#获取word文件
			document = Document();
			document.add_paragraph(text)
			document.save(docdir + "/" + str(id) + '.docx');
		meeting.resultUrl = "/download/txt/" + str(id) + "/";
		meeting.status = "finish";#结束会议
		meeting.wavdate = datetime.datetime.now();#wav上传时间
		meeting.txtdate = datetime.datetime.now();#txt上传时间
		meeting.save();
		result['statusCode'] = 200;
		result["message"] = "ok";
		return Response(result,status=status.HTTP_200_OK);
	result['statusCode'] = 400;
	result["message"] = "error";
	return Response(result,status=status.HTTP_200_OK);

def getMeetingRecord(request, id):#下载获取会议音频
	baseDir = os.path.dirname(os.path.abspath(__name__));
	wavdir = os.path.join(baseDir,'static','wav');
	response = FileResponse(open(wavdir + "/" + str(id) + ".wav", 'rb'))
	return response;

def downloadTxt(request, id):#下载获取会议文本文件
	baseDir = os.path.dirname(os.path.abspath(__name__));
	txtdir = os.path.join(baseDir,'static','txt');
	docdir = os.path.join(baseDir,'static','doc');
	try:
		meeting = Meeting.objects.get(id=id);#查询会议
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
		if meeting.resultType == "txt":#下载txt
			file = io.open(txtdir + "/" + str(id) + ".txt",'r', encoding='UTF-8');
			response = HttpResponse(file);
			response['Content-Type']='application/octet-stream';
			txtname = ((meeting.txtname).encode("utf-8"))
			response['Content-Disposition']='attachment;filename="'+ txtname + '.txt' + '"';
		else :#下载word
			file = io.open(docdir + "/" + str(id) + ".docx",'rb');
			response = HttpResponse(file);
			response['Content-Type']='application/octet-stream';
			txtname = ((meeting.txtname).encode("utf-8"))
			response['Content-Disposition']='attachment;filename="'+ txtname + '.docx' + '"';
		return response
	except Exception, e:
		print e
	return response

def downloadWav(request, id):#下载获取会议音频
	baseDir = os.path.dirname(os.path.abspath(__name__));
	wavdir = os.path.join(baseDir,'static','wav');
	file=open(wavdir + "/" + str(id) + ".wav",'rb'); #打开会议的音频
	response =HttpResponse(file);
	try:
		meeting = Meeting.objects.get(id=id); #查询会议
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
		response['Content-Type']='application/octet-stream';
		wavname = ((meeting.wavname).encode("utf-8")) #获取音频文件名
		response['Content-Disposition']='attachment;filename="'+ (wavname + '.wav' + '"');
		return response
	except Exception, e:
		print e
	return response

def getMeetingObject(meeting): #获取 meeting 的 object格式
	objectMeeting = {};
	objectMeeting['id'] = meeting.id;
	objectMeeting['meetingName'] = meeting.meetingName;
	objectMeeting['description'] = meeting.description;
	objectMeeting['resultType'] = meeting.resultType;
	objectMeeting['createDate'] = meeting.createDate;
	objectMeeting['finishDate'] = meeting.finishDate;
	objectMeeting['userName'] = meeting.userName;
	objectMeeting['status'] = meeting.status;
	objectMeeting['language'] = meeting.language;
	objectMeeting['audioType'] = meeting.audioType;
	objectMeeting['audioBit'] = meeting.audioBit;
	objectMeeting['resultUrl'] = meeting.resultUrl;
	objectMeeting['audioUrl'] = meeting.audioUrl;
	objectMeeting['audioFrequency'] = meeting.audioFrequency;
	return objectMeeting; #返回meeting的object格式

def judgePermission(userName, password):#判断用户权限
	try :
		user = User.objects.get(userName=userName);
		if (user is None): #没找到
			return 404;
		if (user.password != password):
			return 401; #健全错误
		return 200;
	except Exception, e :
		return 404;

@api_view(['POST'])
def openGetAllMeetings(request): #获取全部会议
	result = {};
	requestData = json.loads(request.body);
	userName = requestData['userName'];
	password = requestData['password'];
	statusCode = judgePermission(userName, password); #判断用户权限
	if (statusCode != 200): #用户没有权限
		result['statusCode'] = statusCode;
		result["message"] = "error";
		return Response(result,status=status.HTTP_200_OK);
	meetings = Meeting.objects.all(); #获取全部会议
	meetingList = [];
	for meeting in meetings:
		meetingList.append(getMeetingObject(meeting)); #依次返回
	result["meetings"] = meetingList;
	result['statusCode'] = 200;
	result["message"] = "ok";
	return Response(result,status=status.HTTP_200_OK);

@api_view(['GET'])
def deleteMeetingRecord(request, id): #删除会议音频
	result = {};
	try:
		meeting = Meeting.objects.get(id=id); #查询会议
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
			return Response(result,status=status.HTTP_200_OK); 
	except Exception, e:
		result['statusCode'] = 404;
		result["message"] = "not find";
		return Response(result,status=status.HTTP_200_OK);
	baseDir = os.path.dirname(os.path.abspath(__name__));
	wavdir = os.path.join(baseDir,'static','wav'); #获取wav地址
	my_file = wavdir + "/" + str(id) + ".wav" 
	if os.path.exists(my_file):#删除文件
		os.remove(my_file)
	meeting.audioUrl = "none";
	meeting.save();
	result['statusCode'] = 200;
	result["message"] = "ok";
	return Response(result,status=status.HTTP_200_OK);

@api_view(['GET'])
def deleteMeetingTxt(request, id): #删除文本文件
	result = {};
	try:
		meeting = Meeting.objects.get(id=id); #查询会议
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
			return Response(result,status=status.HTTP_200_OK);
	except Exception, e:
		result['statusCode'] = 404;
		result["message"] = "not find";
		return Response(result,status=status.HTTP_200_OK);
	meeting.resultUrl = "none";
	baseDir = os.path.dirname(os.path.abspath(__name__));
	txtdir = os.path.join(baseDir,'static','txt');#获取文本文件地址
	docdir = os.path.join(baseDir,'static','doc');
	if (meeting.resultType == "txt"): #如果是txt文件
		my_file = txtdir + "/" + str(id) + ".txt"
	else : #如果是word文件
		my_file = docdir + "/" + str(id) + ".docx"
	if os.path.exists(my_file): #删除文件
		os.remove(my_file)
	meeting.save();
	result['statusCode'] = 200;
	result["message"] = "ok";
	return Response(result,status=status.HTTP_200_OK);

@api_view(['POST'])
def openGetUserAllMeetings(request):#获取用户创建的音频
	result = {};
	requestData = json.loads(request.body);
	userName = requestData['userName'];
	password = requestData['password'];
	statusCode = judgePermission(userName, password);#判断用户权限
	if (statusCode != 200):
		result['statusCode'] = statusCode;
		result["message"] = "error";
		return Response(result,status=status.HTTP_200_OK);
	meetings = Meeting.objects.filter(userName=userName);#判断用户名下的音频
	meetingList = [];
	for meeting in meetings:
		meetingList.append(getMeetingObject(meeting));
	result["meetings"] = meetingList;#拼装成rest结构返回
	result['statusCode'] = 200;
	result["message"] = "ok";
	return Response(result,status=status.HTTP_200_OK);

@api_view(['POST'])
def openUpdateMeeting(request): #更新会议信息
	result = {};
	requestData = json.loads(request.body);
	userName = requestData['userName'];
	password = requestData['password'];
	id = requestData['id'];
	meetingName = requestData["meetingName"];
	description = requestData["description"];
	wavname = requestData["wavname"];
	txtname = requestData["txtname"];
	statusCode = judgePermission(userName, password);#判断用户权限
	if (statusCode != 200):
		result['statusCode'] = statusCode;
		result["message"] = "error";
		return Response(result,status=status.HTTP_200_OK);
	try:
		meeting = Meeting.objects.get(id=id); #获取会议
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
			return Response(result,status=status.HTTP_200_OK);
	except Exception, e:
		result['statusCode'] = 404;
		result["message"] = "not find";
		return Response(result,status=status.HTTP_200_OK);
	result["message"] = ""; 
	if meeting.meetingName != meetingName:#依次更改并返回
		meeting.meetingName = meetingName;
		result["message"] = "会议名称更新;";
	if meeting.description != description:
		meeting.description = description;
		result["message"] = result["message"] + "会议描述更新;";
	if meeting.wavname != wavname:
		meeting.wavname = wavname;
		result["message"] = result["message"] + "音频文件名更新;";
	if meeting.txtname != txtname:
		meeting.txtname = txtname;
		result["message"] = result["message"] + "文本文件名更新;";
	meeting.save();
	result['statusCode'] = 200;
	if result["message"] == "":
		result["message"] = "无更新";
	return Response(result,status=status.HTTP_200_OK);


@api_view(['POST'])
def openDeleteMeetings(request):#删除会议
	result = {};
	requestData = json.loads(request.body);
	userName = requestData['userName'];
	password = requestData['password'];
	statusCode = judgePermission(userName, password);#判断权限
	if (statusCode != 200):
		result['statusCode'] = statusCode;
		result["message"] = "error";
		return Response(result,status=status.HTTP_200_OK);
	meetings = Meeting.objects.all();
	meetingList = [];
	for meeting in meetings:
		meetingList.append(getMeetingObject(meeting));#删除后返回
	result["meetings"] = meetingList;
	result['statusCode'] = 200;
	result["message"] = "ok";
	return Response(result,status=status.HTTP_200_OK);

@api_view(['POST'])
def openReadTxt(request):#获取文本信息
	result = {};
	requestData = json.loads(request.body);
	userName = requestData['userName'];
	password = requestData['password'];
	id = requestData['id'];
	statusCode = judgePermission(userName, password);#判断用户权限
	result['statusCode'] = 200;
	result["message"] = "ok";
	baseDir = os.path.dirname(os.path.abspath(__name__));#获取跟路径
	txtdir = os.path.join(baseDir,'static','txt');
	docdir = os.path.join(baseDir,'static','doc');
	
	#根据文件类型初始化文件内容和占地尺寸
	result["txt"] = "";
	result["txtSize"] = 0;
	try:
		meeting = Meeting.objects.get(id=id);
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
			return Response(result,status=status.HTTP_200_OK);
	except Exception, e:
		result['statusCode'] = 404;
		result["message"] = "not find";
		return Response(result,status=status.HTTP_200_OK);
	if meeting.resultType == "txt": #txt文件
		my_file = txtdir + "/" + str(id) + ".txt"
		if os.path.exists(my_file):
			result["txtSize"] = os.path.getsize(my_file)
			f = open(my_file)
			result["txt"] = f.read()
	else : #word文件
		my_file = docdir + "/" + str(id) + ".docx"
		if os.path.exists(my_file):
			result["txtSize"] = os.path.getsize(my_file)
			document = docx.Document(my_file)
			result["txt"] = '\n\n'.join([paragraph.text.encode('utf-8') for paragraph in document.paragraphs])
	return Response(result,status=status.HTTP_200_OK);


@api_view(['POST'])
def openReadWav(request): #获取音频信息
	result = {};
	requestData = json.loads(request.body);
	userName = requestData['userName'];
	password = requestData['password'];
	id = requestData['id'];
	statusCode = judgePermission(userName, password); #判断用户权限
	result['statusCode'] = 200;
	result["message"] = "ok";
	baseDir = os.path.dirname(os.path.abspath(__name__));
	wav = os.path.join(baseDir,'static','wav');
	my_file = wav + "/" + str(id) + ".wav"
	#按照音频具体信息进行赋值
	result['sample_width'] = 0;
	result['channels'] = 0;
	result['frame_rate'] = 0;
	result['duration_seconds'] = 0
	result["wavSize"] = 0
	if os.path.exists(my_file):
		result["wavSize"] = os.path.getsize(my_file)
		wav_audio = AudioSegment.from_file(my_file, format="wav")
		result['sample_width'] = wav_audio.sample_width;
		result['channels'] = wav_audio.channels;
		result['frame_rate'] = wav_audio.frame_rate;
		result['duration_seconds'] = wav_audio.duration_seconds;
	return Response(result,status=status.HTTP_200_OK);

@api_view(['POST'])
def openReplaceTxt(request): #替换txt文件
	result = {};
	requestData = json.loads(request.body);
	id = requestData['id'];
	userName = requestData['userName'];
	password = requestData['password'];
	text = requestData['txt'];
	statusCode = judgePermission(userName, password); #判断用户权限
	result['statusCode'] = 200;
	result["message"] = "ok";
	baseDir = os.path.dirname(os.path.abspath(__name__)); #获取根目录
	txtdir = os.path.join(baseDir,'static','txt');
	docdir = os.path.join(baseDir,'static','doc');
	
	try:
		meeting = Meeting.objects.get(id=id); #获取会议信息
		if (meeting is None):
			result['statusCode'] = 404;
			result["message"] = "not find";
			return Response(result,status=status.HTTP_200_OK);
	except Exception, e:
		result['statusCode'] = 404;
		result["message"] = "not find";
		return Response(result,status=status.HTTP_200_OK);

	if meeting.resultType == "txt": #文本文件
		ftxt = io.open(txtdir + "/" + str(id) + ".txt",'w',encoding='UTF-8');
		ftxt.write(text);
		ftxt.close();
	else : #word文件
		document = Document();
		document.add_paragraph(text)
		document.save(docdir + "/" + str(id) + '.docx');
	meeting.resultUrl = "/download/txt/" + str(id) + "/";
	meeting.save();
	return Response(result,status=status.HTTP_200_OK);

@api_view(['POST'])
def openSearch(request): #搜索音频信息
	result = {};
	requestData = json.loads(request.body);
	userName = requestData['userName'];
	password = requestData['password'];
	meetingName = requestData['meetingName'];
	beginTime = requestData['beginTime'];
	endTime = requestData['endTime'];
	statusCode = judgePermission(userName, password);#判断用户权限
	result['statusCode'] = 200;
	result["message"] = "ok";
	meetings = []

	#根据条件分别进行查询
	if (beginTime == u"" and endTime == u"" and meetingName == u""):
		meetings = Meeting.objects.all();
	elif (beginTime == u"" and endTime == u""):
		meetings = Meeting.objects.filter(meetingName = meetingName)
	elif (beginTime != u"" and endTime == u""):
		beginDate = datetime.datetime.strptime(beginTime, "%Y-%m-%d")
		meetings = Meeting.objects.filter(createDate__gte=beginDate);
	elif (beginTime == u"" and endTime != u""):
		endDate = datetime.datetime.strptime(endTime, "%Y-%m-%d")
		meetings = Meeting.objects.filter(createDate__lte=endDate);
	else :
		beginDate = datetime.datetime.strptime(beginTime, "%Y-%m-%d")
		endData = datetime.datetime.strptime(endTime, "%Y-%m-%d")
		meetings = Meeting.objects.filter(createDate__range=(beginDate, endData));
	meetingList = [];
	for meeting in meetings:
		if (meetingName == u"" or meetingName == meeting.meetingName) and meeting.userName == userName:
			meetingList.append(getMeetingObject(meeting));
	result["meetings"] = meetingList;
	return Response(result,status=status.HTTP_200_OK);
